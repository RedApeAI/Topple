"""Uploaded bytes to text, in memory only.

**The raw document never touches a disk.** Not a temp file, not a cache, not a
bucket. `pypdf` and `python-docx` both accept a file-like object, so the bytes
go into a `BytesIO` and are dropped when this function returns. That is a
deliberate property, not an accident of the API — it is asserted in
`tests/test_document_pipeline.py`, because the easy way to write a PDF parser
is `open(tempfile, "wb")` and nobody would notice.

The other reason parsing lives here rather than in the BFF: the LangGraph
document flow needs the text, and Python's PDF and DOCX libraries are better
than the JavaScript ones. Identity is still resolved at the BFF and forwarded,
so the browser never reaches this service.
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


class UnsupportedDocument(Exception):
    """A file type we cannot read. Maps to HTTP 415."""


class UnreadableDocument(Exception):
    """The right type, but nothing usable came out. Maps to HTTP 422."""


#: Extensions handled natively, with no external process.
SUPPORTED = frozenset(
    {"pdf", "docx", "txt", "md", "markdown", "csv", "tsv", "json", "yaml", "yml", "html", "htm"}
)

#: Still refused, and named individually so the error can say what to do.
UNSUPPORTED_HINTS = {
    "doc": "legacy .doc — open it in Word and save as .docx",
    "pptx": "PowerPoint — export the notes or slides as PDF",
    "xlsx": "Excel — export as CSV",
    "pages": "Apple Pages — export as PDF or Word",
}


def extension_of(filename: str) -> str:
    parts = filename.lower().rsplit(".", 1)
    return parts[1] if len(parts) == 2 else ""


def _strip_html(html: str) -> str:
    import re

    text = re.sub(r"<(script|style)[\s\S]*?</\1>", "", html, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    for entity, char in (("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">")):
        text = text.replace(entity, char)
    return text


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    # BytesIO, not a temp file — see the module docstring.
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # An empty-password decrypt covers the common "protected but not
        # secret" case; a real password is the user's to remove.
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise UnreadableDocument(
                "That PDF is password protected. Remove the password and try again."
            ) from exc

    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 — one bad page must not lose the rest
            logger.warning("pdf page %d unreadable: %s", number, exc)
    return "\n\n".join(p for p in pages if p.strip())


def _read_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables carry the numbers in most commercial documents — a price sheet is
    # usually a table, and dropping them would leave a knowledge base of
    # headings and prose with no figures in it.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Text from an uploaded file. Raises rather than returning junk."""
    extension = extension_of(filename)

    if extension in UNSUPPORTED_HINTS:
        raise UnsupportedDocument(
            f"Can't read {filename} — {UNSUPPORTED_HINTS[extension]}."
        )
    if extension not in SUPPORTED:
        raise UnsupportedDocument(
            f"Can't read a .{extension or '?'} file. "
            f"Supported: {', '.join(sorted(SUPPORTED))}."
        )

    if extension == "pdf":
        text = _read_pdf(data)
        if not text.strip():
            # A scanned PDF is images; pypdf finds no text layer. Saying so
            # beats storing an empty document and wondering later.
            raise UnreadableDocument(
                "No text layer in that PDF — it looks like a scan. "
                "Run it through OCR first, or upload a text version."
            )
        return text

    if extension == "docx":
        text = _read_docx(data)
        if not text.strip():
            raise UnreadableDocument("That Word document has no readable text.")
        return text

    decoded = data.decode("utf-8", errors="replace")
    if extension in ("html", "htm"):
        decoded = _strip_html(decoded)

    # A binary renamed to .txt decodes to replacement characters; embedding it
    # would fill the index with noise that still matches queries.
    if decoded.count("�") > max(len(decoded) * 0.02, 8):
        raise UnsupportedDocument(
            "That file doesn't appear to be text. Check the format and try again."
        )
    if not decoded.strip():
        raise UnreadableDocument("There was no readable text in that file.")
    return decoded
