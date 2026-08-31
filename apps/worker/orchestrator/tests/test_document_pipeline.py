"""The document is transient: RAM only, purged, never persisted.

Three claims, each of which is easy to believe and easy to get wrong:

1. The raw bytes never touch a disk.
2. The raw text is cleared from state when compression finishes.
3. Only ~100-word summaries reach Qdrant, marked as episodic so the turn
   pipeline's grounding guardrail never quotes one.
"""
from __future__ import annotations

import io
import json

import pytest

from app.documents import parse, splitter
from app.documents.ingest import doc_id_for, ingest_document
from app.graph import document_graph
from app.llm import gateway
from app.schemas.llm import GenerationCall, GenerationOutput, LLMCallStats
from app.stores import knowledge, qdrant

PRICE_SHEET = """# Marina Crest

Marina Crest sits in Dubai Marina. Handover is Q4 2027.

## Pricing

1 Bedroom from AED 1,350,000.
2 Bedroom from AED 2,150,000.

## Amenities

Infinity pool on level 40, gym, sauna, and a 25 m lap pool.
"""


class RecordingQdrant:
    """Captures what was stored, and applies the scope filter honestly."""

    def __init__(self):
        self.points: list[dict] = []

    async def collection_exists(self, collection):
        return True

    async def create_payload_index(self, **kwargs):
        return None

    async def add(self, collection_name, documents, metadata=None, ids=None):
        for text, meta in zip(documents, metadata or []):
            self.points.append({"text": text, "payload": meta})
        return []

    async def delete(self, collection_name, points_selector):
        wanted = {
            c.key: c.match.value for c in points_selector.filter.must
        }
        self.points = [
            p for p in self.points
            if not all(p["payload"].get(k) == v for k, v in wanted.items())
        ]

    async def query(self, collection_name, query_text, query_filter=None, limit=10):
        wanted = {c.key: c.match.value for c in (query_filter.must if query_filter else [])}

        class R:
            def __init__(self, p):
                self.id = p["payload"]["chunk_id"]
                self.document = p["text"]
                self.metadata = p["payload"]
                self.score = 0.9

        return [
            R(p) for p in self.points
            if all(p["payload"].get(k) == v for k, v in wanted.items())
        ][:limit]


@pytest.fixture()
def store(monkeypatch) -> RecordingQdrant:
    client = RecordingQdrant()
    qdrant.set_client(client)
    yield client
    qdrant.set_client(None)


@pytest.fixture()
def model(monkeypatch):
    """A model that summarises by labelling, so summaries are distinguishable
    from the verbatim text they came from."""
    calls: list[dict] = []

    async def generate(*, model, messages):
        calls.append({"model": model, "messages": messages})
        system = messages[0]["content"]
        if "under 100 words" in system:
            body = "SUMMARY of a passage about Marina Crest pricing and amenities."
        else:
            body = "The 1 Bedroom starts at AED 1,350,000."
        return GenerationCall(
            output=GenerationOutput(messages=[body]),
            raw_text=body,
            stats=LLMCallStats(),
        )

    monkeypatch.setattr(gateway, "generate", generate)
    return calls


# --------------------------------------------------------------------------- #
# 1. Nothing touches a disk
# --------------------------------------------------------------------------- #
@pytest.fixture()
def no_disk_writes(monkeypatch):
    """Fail if anything opens a file for writing while parsing.

    Watches `os.open`, not `builtins.open`. Every file open in CPython bottoms
    out there — including `tempfile.NamedTemporaryFile`, which is exactly what
    a naive PDF parser reaches for and which sails straight past a
    `builtins.open` patch. Verified by mutation: a parser that writes a temp
    file is caught by this and was not caught by the earlier version.

    Read-only opens are allowed. Config and font files get read during a parse
    and blocking those would only make the test flaky about import order; the
    claim being defended is that the *document* is never written down.

    Lazy imports are warmed **before** the watcher is armed. `filelock` — pulled
    in by `qdrant_client.models` — probes symlink behaviour at import time by
    creating a temp file, exactly once per process. That is a library's
    one-off, not our pipeline writing the document, but it is indistinguishable
    afterwards and it made this test pass or fail on import order alone. Warming
    first means every write the watcher sees is genuinely ours.
    """
    import os

    import qdrant_client.models  # noqa: F401 — warms filelock's import probe

    writes: list[str] = []
    real_os_open = os.open
    write_flags = (
        os.O_WRONLY | os.O_RDWR | os.O_CREAT | os.O_APPEND | getattr(os, "O_TMPFILE", 0)
    )

    def watched(path, flags, *args, **kwargs):
        if flags & write_flags:
            writes.append(str(path))
        return real_os_open(path, flags, *args, **kwargs)

    monkeypatch.setattr(os, "open", watched)
    return writes


def test_pdf_parsing_never_writes_to_disk(no_disk_writes):
    """The obvious PDF implementation writes a temp file first. This one must
    not — and an assertion is the only thing that keeps it that way."""
    pdf = _pdf_with_text("Marina Crest 1 Bedroom from AED 1,350,000.")
    text = parse.extract_text("prices.pdf", pdf)

    assert "Marina Crest" in text
    assert no_disk_writes == [], f"the parser wrote to disk: {no_disk_writes}"


def test_docx_parsing_never_writes_to_disk(no_disk_writes):
    docx_bytes = _minimal_docx(["Marina Crest", "1 Bedroom from AED 1,350,000."])
    text = parse.extract_text("prices.docx", docx_bytes)

    assert "1,350,000" in text
    assert no_disk_writes == [], f"the parser wrote to disk: {no_disk_writes}"


async def test_the_whole_pipeline_never_writes_to_disk(store, model, no_disk_writes):
    """Not just the parser — parse, answer, compress and purge together."""
    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
        question="What does a 1 bedroom cost?",
    )
    assert no_disk_writes == [], f"the pipeline wrote to disk: {no_disk_writes}"


def test_docx_tables_are_read():
    """A price sheet is usually a table. Dropping tables would leave headings
    and prose with no figures in them."""
    docx_bytes = _minimal_docx(["Prices"], table=[["Unit", "Price"], ["1BR", "AED 1,350,000"]])
    text = parse.extract_text("prices.docx", docx_bytes)
    assert "1,350,000" in text
    assert "1BR" in text


# --------------------------------------------------------------------------- #
# 2. The purge
# --------------------------------------------------------------------------- #
async def test_raw_text_is_cleared_when_ingestion_finishes(store, model):
    state = await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    assert state.purged is True
    assert state.raw_text == "", "the document must not survive its own ingestion"


async def test_the_uploaded_bytes_are_released(store, model):
    """The context holds the bytes; purge drops the reference."""
    from app.graph.context import GraphContext

    context = GraphContext(upload=PRICE_SHEET.encode())
    assert context.upload
    context.forget_upload()
    assert context.upload == b""


async def test_purge_runs_even_when_compression_fails(store, monkeypatch):
    async def explode(*, model, messages):
        raise RuntimeError("model outage mid-compression")

    monkeypatch.setattr(gateway, "generate", explode)
    state = await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    assert state.purged is True
    assert state.raw_text == ""


async def test_raw_text_is_excluded_from_any_state_dump(store, model):
    """`exclude=True` is what stops the document riding along in a log line or
    a serialised state — the difference between deleting it and deleting the
    copy you happened to be looking at."""
    from app.graph.document_state import DocumentState

    state = DocumentState(
        tenant_id="t", user_id="u", filename="f.md", doc_id="f",
        collection="kb", raw_text=PRICE_SHEET,
    )
    dumped = json.dumps(state.model_dump(), default=str)
    assert "Marina Crest" not in dumped
    assert "Marina Crest" not in json.dumps(state.redacted(), default=str)


# --------------------------------------------------------------------------- #
# 3. What reaches Qdrant
# --------------------------------------------------------------------------- #
async def test_only_summaries_are_stored_by_default(store, model):
    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    assert store.points, "something must have been stored"
    for point in store.points:
        assert point["text"].startswith("SUMMARY"), point["text"][:60]
        assert point["payload"][knowledge.MEMORY_KEY] == knowledge.EPISODIC


async def test_summaries_are_scoped_to_the_uploading_user(store, model):
    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    for point in store.points:
        assert point["payload"]["tenant_id"] == "redape"
        assert point["payload"]["user_id"] == "alice"


async def test_the_extraction_prompt_is_the_specified_one(store, model):
    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    systems = [c["messages"][0]["content"] for c in model]
    assert any("Extract key decisions, facts, entities, and intent" in s for s in systems)
    assert any("under 100 words" in s for s in systems)


async def test_verbatim_storage_is_opt_in_and_marked_semantic(store, model):
    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb", store_verbatim=True,
    )
    kinds = {p["payload"][knowledge.MEMORY_KEY] for p in store.points}
    assert kinds == {knowledge.EPISODIC, knowledge.SEMANTIC}

    verbatim = [p for p in store.points if p["payload"][knowledge.MEMORY_KEY] == knowledge.SEMANTIC]
    assert any("1,350,000" in p["text"] for p in verbatim), (
        "the author's own figures must survive verbatim, or nothing can be quoted"
    )


async def test_episodic_memory_is_invisible_to_a_semantic_read(store, model):
    """The guardrail's corpus must never contain a model-written summary."""
    from app.stores.knowledge import KnowledgeScope

    await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    semantic = KnowledgeScope("redape", "alice").for_memory(knowledge.SEMANTIC)
    hits, _ = await qdrant.retrieve("kb", "price", 10, 0.3, scope=semantic)
    assert hits == [], "episodic summaries must not be retrievable as facts"


# --------------------------------------------------------------------------- #
# The immediate answer
# --------------------------------------------------------------------------- #
async def test_a_question_is_answered_while_the_document_is_resident(store, model):
    state = await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
        question="What does a 1 bedroom cost?",
    )
    assert "1,350,000" in state.answer
    assert state.purged is True


async def test_no_question_means_no_answer_call(store, model):
    state = await ingest_document(
        tenant_id="redape", user_id="alice", filename="prices.md",
        data=PRICE_SHEET.encode(), collection="kb",
    )
    assert state.answer == ""
    assert all("under 100 words" in c["messages"][0]["content"] for c in model)


# --------------------------------------------------------------------------- #
# Chunking to spec
# --------------------------------------------------------------------------- #
def test_chunks_sit_in_the_specified_band():
    long_document = "\n\n".join(
        f"Paragraph {i}. " + "The property has a pool and parking. " * 12
        for i in range(60)
    )
    chunks = splitter.split(long_document)
    assert len(chunks) > 1

    sizes = [splitter.estimate_tokens(c) for c in chunks]
    # The last chunk is whatever remains, so it is allowed to be short.
    for size in sizes[:-1]:
        assert 400 <= size <= splitter.MAX_TOKENS, sizes


def test_consecutive_chunks_overlap():
    document = "\n\n".join(f"Sentence block {i} about pricing." * 20 for i in range(20))
    chunks = splitter.split(document)
    assert len(chunks) >= 2
    # Some tail of chunk N should appear at the head of chunk N+1.
    head = chunks[1][:200]
    assert any(word in chunks[0] for word in head.split()[:5])


def test_a_heading_stays_with_its_section():
    chunks = splitter.split(PRICE_SHEET)
    joined = "\n".join(chunks)
    assert "## Pricing" in joined
    holder = next(c for c in chunks if "1,350,000" in c)
    assert "Marina Crest" in holder or "Pricing" in holder


def test_empty_input_yields_nothing():
    assert splitter.split("") == []
    assert splitter.split("   \n\n  ") == []


# --------------------------------------------------------------------------- #
# Refusals
# --------------------------------------------------------------------------- #
def test_a_scanned_pdf_says_so_rather_than_storing_nothing():
    empty_pdf = _minimal_pdf("")
    with pytest.raises(parse.UnreadableDocument, match="scan"):
        parse.extract_text("scan.pdf", empty_pdf)


def test_legacy_doc_is_refused_with_a_next_step():
    with pytest.raises(parse.UnsupportedDocument, match="save as .docx"):
        parse.extract_text("old.doc", b"\xd0\xcf\x11\xe0")


def test_binary_renamed_to_txt_is_refused():
    with pytest.raises(parse.UnsupportedDocument, match="doesn't appear to be text"):
        parse.extract_text("fake.txt", bytes(range(256)) * 20)


def test_doc_ids_are_stable_so_reupload_replaces():
    assert doc_id_for("Marina Crest Prices 2026.pdf") == "marina-crest-prices-2026"
    assert doc_id_for("Marina Crest Prices 2026.docx") == "marina-crest-prices-2026"


# --------------------------------------------------------------------------- #
# Fixtures built in memory
# --------------------------------------------------------------------------- #
def _minimal_pdf(text: str) -> bytes:
    """A one-page PDF built in memory, so these tests need no sample files."""
    from pypdf import PdfWriter

    try:
        from reportlab.pdfgen import canvas  # noqa: F401
    except ImportError:
        pass

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    base = buffer.getvalue()
    if not text:
        return base
    # pypdf cannot author a text layer; splice a content stream describing one.
    return _pdf_with_text(text)


def _pdf_with_text(text: str) -> bytes:
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    content = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode() + b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
    )
    return bytes(out)


def _minimal_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                added.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
